from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/PowerBI_PPU_API_Access_Request_For_Admin.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    para = doc.add_paragraph(style=f"Heading {level}")
    run = para.add_run(text)
    set_run(run, bold=True)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(text)
    return para


def add_number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.add_run(text)
    return para


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Power BI PPU Monitoring API Access Request").font.size = Pt(9)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], "F2F4F7")
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                set_run(run, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_width(table, widths)
    return table


def build_doc():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Power BI Premium Per User Monitoring API Access Request")

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        "Request for read-only API access to build an internal Power BI PPU governance, usage, and refresh monitoring dashboard."
    )
    set_run(subtitle_run, color="44546A")

    add_heading(doc, "Purpose", 1)
    doc.add_paragraph(
        "We are building an internal monitoring dashboard for Power BI Premium Per User (PPU). "
        "The dashboard will help IT and BI governance teams monitor workspace inventory, report usage, dataset refresh health, idle reports, ownership, and governance risk."
    )
    doc.add_paragraph(
        "We are not requesting Microsoft Fabric capacity monitoring, dedicated Premium capacity CPU/memory metrics, or write access to Power BI content. This request is for read-only Power BI monitoring APIs."
    )

    add_heading(doc, "Access Requested From Admin", 1)
    add_number(doc, "Create or approve a Microsoft Entra app registration for the monitoring connector.")
    add_number(doc, "Create a dedicated Microsoft Entra security group for the app, for example: PowerBI-Monitoring-API-Access.")
    add_number(doc, "Add the app/service principal to the security group.")
    add_number(doc, "Enable read-only Power BI admin API access for that security group in the Power BI Admin Portal.")
    add_number(doc, "Provide the connection details listed in the checklist below through the approved secure secret-sharing process.")

    add_heading(doc, "Tenant Settings To Enable", 1)
    add_table(
        doc,
        ["Setting", "Required?", "Reason"],
        [
            [
                "Service principals can access read-only admin APIs",
                "Yes",
                "Allows the monitoring app to read tenant-level workspace, report, dataset, and activity metadata.",
            ],
            [
                "Enhance admin APIs responses with detailed metadata",
                "Yes",
                "Returns richer metadata for governance and ownership monitoring.",
            ],
            [
                "Enhance admin APIs responses with DAX and mashup expressions",
                "Optional",
                "Only needed if admins want deeper semantic model governance later. Not required for the first version.",
            ],
            [
                "Allow service principals to use Power BI APIs",
                "Yes, if required by tenant configuration",
                "Allows the service principal to authenticate to Power BI APIs.",
            ],
        ],
        [2900, 1500, 4960],
    )

    add_heading(doc, "API Endpoints Required", 1)
    add_table(
        doc,
        ["Use", "Endpoint", "Access Type"],
        [
            [
                "Workspace, report, dataset, dashboard, and user inventory",
                "GET https://api.powerbi.com/v1.0/myorg/admin/groups?$top=5000&$expand=reports,datasets,dashboards,users",
                "Power BI read-only admin API",
            ],
            [
                "Report usage and activity events",
                "GET https://api.powerbi.com/v1.0/myorg/admin/activityevents",
                "Power BI read-only admin API",
            ],
            [
                "Dataset refresh history",
                "GET https://api.powerbi.com/v1.0/myorg/groups/{workspaceId}/datasets/{datasetId}/refreshes",
                "Power BI REST API",
            ],
            [
                "Dataset/semantic model list per workspace",
                "GET https://api.powerbi.com/v1.0/myorg/groups/{workspaceId}/datasets",
                "Power BI REST API",
            ],
            [
                "Report list per workspace",
                "GET https://api.powerbi.com/v1.0/myorg/groups/{workspaceId}/reports",
                "Power BI REST API",
            ],
            [
                "Dashboard list per workspace",
                "GET https://api.powerbi.com/v1.0/myorg/groups/{workspaceId}/dashboards",
                "Power BI REST API",
            ],
            [
                "User department, job title, and active/inactive status",
                "GET https://graph.microsoft.com/v1.0/users?$select=id,displayName,userPrincipalName,mail,department,jobTitle,accountEnabled",
                "Optional Microsoft Graph application permission",
            ],
        ],
        [2050, 5210, 2100],
    )

    add_heading(doc, "Connection Details Needed From Admin", 1)
    add_table(
        doc,
        ["Item", "Value To Provide", "Notes"],
        [
            ["Tenant ID", "", "Microsoft Entra tenant/directory ID."],
            ["Client ID", "", "Application/client ID of the app registration."],
            ["Client secret or certificate", "", "Provide through secure secret manager only. Do not send in plain email."],
            ["Service principal object ID", "", "Useful for security group and audit validation."],
            ["Security group name/object ID", "", "Group enabled in Power BI Admin Portal tenant settings."],
            ["Approved scope", "", "All tenant workspaces or a specific list of workspaces/domains."],
            ["Graph approval", "", "Confirm whether Microsoft Graph user details are approved."],
        ],
        [2300, 2600, 4460],
    )

    add_heading(doc, "Authentication Method", 1)
    doc.add_paragraph("The connector will use OAuth 2.0 client credentials.")
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Token URL", "POST https://login.microsoftonline.com/{tenantId}/oauth2/v2.0/token"],
            ["Power BI scope", "https://analysis.windows.net/powerbi/api/.default"],
            ["Auth header for API calls", "Authorization: Bearer {access_token}"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "Security Notes", 1)
    add_bullet(doc, "This request is read-only. No report publishing, dataset editing, workspace deletion, or user permission changes are requested.")
    add_bullet(doc, "Do not send client secrets in email or chat. Use the approved company secret-sharing process.")
    add_bullet(doc, "A certificate is preferred over a client secret if the organization has certificate-based app authentication standards.")
    add_bullet(doc, "The app should be limited to the dedicated monitoring security group rather than enabled for the entire organization where possible.")
    add_bullet(doc, "For Microsoft Graph user enrichment, approve only the minimum required permission such as User.Read.All or Directory.Read.All, based on internal policy.")

    add_heading(doc, "What The Dashboard Will Show", 1)
    add_bullet(doc, "Workspace inventory and ownership.")
    add_bullet(doc, "Report and dashboard inventory.")
    add_bullet(doc, "Dataset/semantic model inventory.")
    add_bullet(doc, "Refresh success, failure, duration, and error details.")
    add_bullet(doc, "Report usage and active users from activity events.")
    add_bullet(doc, "Idle reports, stale assets, and governance recommendations.")

    add_heading(doc, "Important PPU Clarification", 1)
    doc.add_paragraph(
        "Because we use Power BI Premium Per User, this request does not include dedicated capacity CPU, memory, overload, throttling, or Fabric Capacity Metrics. "
        "The first monitoring version will focus on API-supported PPU governance signals: usage, refreshes, workspaces, reports, datasets, owners, and users."
    )

    add_heading(doc, "Microsoft Reference Links", 1)
    refs = [
        ("Power BI Admin - Get Groups As Admin", "https://learn.microsoft.com/en-us/rest/api/power-bi/admin/groups-get-groups-as-admin"),
        ("Power BI Admin - Get Activity Events", "https://learn.microsoft.com/en-us/rest/api/power-bi/admin/get-activity-events"),
        ("Power BI Dataset - Get Refresh History In Group", "https://learn.microsoft.com/en-us/rest/api/power-bi/datasets/get-refresh-history-in-group"),
        ("Power BI service principal authentication", "https://learn.microsoft.com/en-us/power-bi/developer/embedded/embed-service-principal"),
        ("Microsoft Graph - List users", "https://learn.microsoft.com/en-us/graph/api/user-list"),
    ]
    for label, url in refs:
        para = doc.add_paragraph(style="List Bullet")
        add_hyperlink(para, label, url)

    add_heading(doc, "Approval Checklist", 1)
    checklist_items = [
        "Entra app registration created or approved.",
        "Service principal added to dedicated security group.",
        "Power BI read-only admin API tenant setting enabled for the security group.",
        "Detailed metadata admin API response enabled.",
        "Tenant ID and Client ID provided.",
        "Client secret/certificate shared securely.",
        "Graph user enrichment approved or rejected.",
        "Workspace scope confirmed.",
    ]
    for item in checklist_items:
        add_bullet(doc, f"[ ] {item}")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
